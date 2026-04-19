"""Génération de fichiers : DOCX, PDF, XLSX."""

import io
import uuid
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from services import storage_service


def build_cv_docx(
    personal: Dict[str, Any],
    experience: List[Dict[str, Any]],
    education: List[Dict[str, Any]],
    skills: List[str],
    language: str,
) -> bytes:
    """Construit un CV au format Word."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    name = personal.get("full_name") or "CV"
    doc.add_heading(name, 0)
    if personal.get("email"):
        doc.add_paragraph(personal["email"])
    if personal.get("phone"):
        doc.add_paragraph(personal["phone"])
    if personal.get("location"):
        doc.add_paragraph(personal["location"])
    if personal.get("summary"):
        doc.add_heading("Profil" if language.startswith("fr") else "Profile", level=1)
        doc.add_paragraph(personal["summary"])

    doc.add_heading("Expérience" if language.startswith("fr") else "Experience", level=1)
    for exp in experience:
        p = doc.add_paragraph()
        p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}").bold = True
        doc.add_paragraph(f"{exp.get('start', '')} – {exp.get('end', '')}")
        if exp.get("description"):
            doc.add_paragraph(exp["description"])

    doc.add_heading("Formation" if language.startswith("fr") else "Education", level=1)
    for ed in education:
        doc.add_paragraph(
            f"{ed.get('degree', '')} — {ed.get('school', '')} ({ed.get('year', '')})",
        )

    if skills:
        doc.add_heading("Compétences" if language.startswith("fr") else "Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_letter_docx(
    content: str,
    title: str = "Lettre",
) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_report_pdf(title: str, sections: List[str], body: str) -> bytes:
    """PDF simple avec reportlab."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, title[:80])
    y -= 30
    c.setFont("Helvetica", 11)
    for sec in sections:
        if y < 80:
            c.showPage()
            y = height - 50
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, sec[:100])
        y -= 18
        c.setFont("Helvetica", 11)
    for line in body.split("\n"):
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:120])
        y -= 14
    c.save()
    return buf.getvalue()


def build_excel_workbook(
    title: str,
    columns: List[str],
    rows: List[List[Any]],
) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31] if title else "Sheet"
    ws.append(columns)
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


async def upload_generated(
    data: bytes,
    folder: str,
    filename: str,
    content_type: str,
) -> Dict[str, str]:
    key, url = await storage_service.upload_bytes(
        data,
        f"generated/{folder}",
        filename,
        content_type,
    )
    return {"object_key": key, "url": url, "filename": filename}
